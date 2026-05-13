# utils.py — Helper functions (PDF, DOCX, save, tokens)
import os, re, tempfile
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from fpdf import FPDF
import streamlit as st
from config import INPUT_COST_PER_1M, OUTPUT_COST_PER_1M

# ── Multi-format Resume Extraction (PDF, MD, TXT, TYP) ──
@st.cache_data
def extract_resume_text(uploaded_file):
    """Extract text from PDF, Markdown, Typst, or plain text files."""
    fname = uploaded_file.name.lower()
    content = uploaded_file.getvalue()

    if fname.endswith('.pdf'):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        text = ""
        try:
            # Try 1: PyPDFLoader via LangChain
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(tmp_path)
            docs = loader.load()
            text = "\n\n".join(doc.page_content for doc in docs)
        except Exception:
            pass

        if not text.strip():
            try:
                # Try 2: pypdf directly
                from pypdf import PdfReader
                reader = PdfReader(tmp_path)
                pages = []
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
                text = "\n\n".join(pages)
            except Exception:
                pass

        if not text.strip():
            try:
                # Try 3: Read raw PDF bytes for text content
                raw = content.decode('latin-1', errors='ignore')
                # Extract anything between stream/endstream that looks like text
                import re
                streams = re.findall(r'\((.*?)\)', raw)
                if streams:
                    text = "\n".join(s for s in streams if len(s) > 3 and any(c.isalpha() for c in s))
            except Exception:
                pass

        os.unlink(tmp_path)

        if not text.strip():
            return "[ERROR: Could not extract text from this PDF. It may be image-based. Use .md or .typ format instead.]"
        return text

    elif fname.endswith(('.md', '.txt', '.typ', '.markdown')):
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            return content.decode('latin-1', errors='replace')

    else:
        return f"[Unsupported format: {fname}]"

def is_valid_resume_text(text):
    """Check if extracted text is actual resume content (not an error)."""
    if not text or not text.strip():
        return False
    if text.startswith("[ERROR:") or text.startswith("[Could not") or text.startswith("[Unsupported"):
        return False
    return len(text.strip()) > 50

# ── DOCX Generation (proper formatting) ──
def generate_docx(text, title="Document"):
    """Generate a general-purpose .docx with clean formatting."""
    doc = Document()
    _setup_doc_margins(doc)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        elif stripped.startswith('### '):
            h = doc.add_heading(stripped[4:], level=3)
            _set_heading_font(h, 'Calibri')
        elif stripped.startswith('## '):
            h = doc.add_heading(stripped[3:], level=2)
            _set_heading_font(h, 'Calibri')
        elif stripped.startswith('# '):
            h = doc.add_heading(stripped[2:], level=1)
            _set_heading_font(h, 'Calibri')
        elif stripped.startswith('**') and stripped.endswith('**'):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped.strip('**'))
            run.bold = True
        elif stripped.startswith('\u2022 ') or stripped.startswith('- '):
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(3)
        elif stripped.startswith('| '):
            p = doc.add_paragraph(stripped)
            p.paragraph_format.line_spacing = 1.0
            run = p.runs[0] if p.runs else p.add_run('')
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        else:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generate_docx_cover_letter(text):
    """Generate a properly formatted cover letter .docx.
    Format: Title centered (14pt bold), body justified (11pt, 1.15 spacing),
    date + signature at bottom."""
    doc = Document()
    _setup_doc_margins(doc)

    # Normal style for body
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = [l.strip() for l in text.split('\n') if l.strip()]

    # Filter out markdown headers, "Cover Letter" title etc.
    body_lines = []
    title_line = "Cover Letter"
    for line in lines:
        if line.startswith('#'):
            continue
        if line.lower().startswith('cover letter') and len(line) < 30:
            title_line = line.strip('#').strip()
            continue
        body_lines.append(line)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = 1  # CENTER
    title_p.paragraph_format.space_after = Pt(10)
    title_run = title_p.add_run(title_line)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = 'Calibri'

    # Body paragraphs
    sig_mode = None  # None, 'first', 'rest'
    for line in body_lines:
        if not line:
            continue

        # Detect signature block start
        if line.lower().startswith('sincerely') and len(line) < 20:
            sig_mode = 'first'

        p = doc.add_paragraph()
        p.alignment = 3  # JUSTIFY

        if sig_mode == 'first':
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            sig_mode = 'rest'
        elif sig_mode == 'rest':
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
        else:
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(10)

        run = p.add_run(line)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # Date at the bottom (only if not already present in signature)
    has_date = any('date:' in l.lower() for l in body_lines[-4:])
    if not has_date:
        from datetime import datetime
        date_p = doc.add_paragraph()
        date_p.alignment = 3
        date_p.paragraph_format.space_before = Pt(20)
        date_p.paragraph_format.space_after = Pt(0)
        date_run = date_p.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.name = 'Calibri'
        date_run.font.size = Pt(11)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def _setup_doc_margins(doc):
    """Set standard 1-inch margins on all sides."""
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

def _set_heading_font(heading, font_name='Calibri'):
    """Ensure heading uses the specified font."""
    for run in heading.runs:
        run.font.name = font_name

# ── CV DOCX Template Generator (matches Sarangi CV format) ──
def generate_docx_cv_template(cv_data):
    """Generate a CV .docx matching the Sarangi template format exactly."""
    doc = Document()

    # Margins: match template (1.143" left/right, 0.914" top/bottom)
    for section in doc.sections:
        section.top_margin = Inches(0.914)
        section.bottom_margin = Inches(0.914)
        section.left_margin = Inches(1.143)
        section.right_margin = Inches(1.143)

    # ── 1. NAME (Heading 1) ──
    h = doc.add_heading(cv_data.get("name", ""), level=1)
    _set_heading_font(h, 'Calibri')

    # ── 2. CONTACT (First Paragraph style, bold labels) ──
    contact = cv_data.get("contact", {})
    contact_parts = []
    if contact.get("location"): contact_parts.append(("Location: ", contact["location"]))
    if contact.get("phone"): contact_parts.append(("Phone: ", contact["phone"]))
    if contact.get("email"): contact_parts.append(("Email: ", contact["email"]))
    if contact.get("linkedin"): contact_parts.append(("LinkedIn: ", contact["linkedin"]))
    if contact.get("github"): contact_parts.append(("GitHub: ", contact["github"]))
    if contact.get("scholar"): contact_parts.append(("Google Scholar: ", contact["scholar"]))

    p_contact = doc.add_paragraph()
    p_contact.style = doc.styles['Normal']
    for i, (label, value) in enumerate(contact_parts):
        if i > 0:
            p_contact.add_run(" | ")
        run_label = p_contact.add_run(label)
        run_label.bold = True
        run_label.font.name = 'Calibri'
        run_label.font.size = Pt(10)
        run_value = p_contact.add_run(value)
        run_value.font.name = 'Calibri'
        run_value.font.size = Pt(10)
    p_contact.paragraph_format.space_after = Pt(4)

    # ── Helper: add section ──
    def add_heading2(text):
        h2 = doc.add_heading(text, level=2)
        _set_heading_font(h2, 'Calibri')
        h2.paragraph_format.space_before = Pt(10)

    def add_heading3(text):
        h3 = doc.add_heading(text, level=3)
        _set_heading_font(h3, 'Calibri')
        h3.paragraph_format.space_before = Pt(6)
        h3.paragraph_format.space_after = Pt(2)

    def add_body(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.bold = bold
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.line_spacing = 1.0
        if bold_prefix:
            run_b = p.add_run(bold_prefix + " ")
            run_b.bold = True
            run_b.font.name = 'Calibri'
            run_b.font.size = Pt(10)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        return p

    # ── 3. PROFESSIONAL SUMMARY ──
    add_heading2("Professional Summary")
    add_body(cv_data.get("summary", ""))

    # ── 4. TECHNICAL SKILLS ──
    add_heading2("Technical Skills")
    for skill_cat in cv_data.get("skills", []):
        add_heading3(skill_cat.get("category", ""))
        add_body(skill_cat.get("items", ""))

    # ── 5. WORK EXPERIENCE ──
    add_heading2("Work Experience")
    for exp in cv_data.get("experience", []):
        add_heading3(exp.get("title", ""))
        company_line = f"{exp.get('company', '')} | {exp.get('location', '')} | {exp.get('dates', '')}"
        add_body(company_line, bold=True)

        if exp.get("project_name"):
            add_body(exp["project_name"])

        for bullet in exp.get("highlights", []):
            add_bullet(bullet)

        if exp.get("technologies"):
            add_body("Technologies: " + exp["technologies"], bold=True)

    # ── 6. PUBLICATIONS ──
    pubs = cv_data.get("publications", [])
    if pubs:
        add_heading2("Journal Publications")
        for pub in pubs:
            add_bullet(pub)

    # ── 7. EDUCATION ──
    edu = cv_data.get("education", [])
    if edu:
        add_heading2("Education")
        for e in edu:
            line = f"{e.get('degree', '')}, {e.get('school', '')}, {e.get('year', '')}"
            add_bullet(line, bold_prefix=e.get('degree', ''))

    # ── 8. CERTIFICATIONS ──
    certs = cv_data.get("certifications", [])
    if certs:
        add_heading2("Certifications")
        for cert in certs:
            add_bullet(cert.get("title", ""), bold_prefix=cert.get("title", ""))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ── PDF Generation ──
def generate_pdf(text, title="Document"):
    """Convert text to PDF. Handles special characters and long lines."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", size=11)

    # Add DejaVu as fallback for Unicode symbols
    try:
        pdf.add_font("DejaVu", "", r"C:\Windows\Fonts\DejaVuSans.ttf", uni=True)
        pdf.add_font("DejaVu", "B", r"C:\Windows\Fonts\DejaVuSans-Bold.ttf", uni=True)
    except Exception:
        pass

    for line in text.split('\n'):
        stripped = _sanitize(line.strip())
        if not stripped:
            pdf.cell(0, 6, "", ln=True)
            continue
        try:
            if stripped.startswith('### '):
                pdf.set_font("Helvetica", 'B', 12)
                pdf.multi_cell(0, 8, stripped[4:])
                pdf.ln(1)
                pdf.set_font("Helvetica", size=11)
            elif stripped.startswith('## '):
                pdf.set_font("Helvetica", 'B', 14)
                pdf.multi_cell(0, 10, stripped[3:])
                pdf.ln(1)
                pdf.set_font("Helvetica", size=11)
            elif stripped.startswith('# '):
                pdf.set_font("Helvetica", 'B', 16)
                pdf.multi_cell(0, 12, stripped[2:])
                pdf.ln(1)
                pdf.set_font("Helvetica", size=11)
            elif stripped.startswith('**') and stripped.endswith('**'):
                pdf.set_font("Helvetica", 'B', 11)
                pdf.multi_cell(0, 6, stripped.strip('**'))
                pdf.set_font("Helvetica", size=11)
            elif stripped.startswith('\u2022 ') or stripped.startswith('- '):
                pdf.cell(8, 6, "-", ln=False)
                pdf.multi_cell(0, 6, stripped[2:])
            elif stripped.startswith('| '):
                pdf.set_font("Courier", size=9)
                pdf.cell(0, 6, stripped[:140], ln=True)
                pdf.set_font("Helvetica", size=11)
            else:
                pdf.multi_cell(0, 6, stripped)
        except Exception:
            # Fallback: render as plain text with truncation
            pdf.set_font("Helvetica", size=9)
            try:
                pdf.multi_cell(0, 5, stripped[:200])
            except Exception:
                pdf.cell(0, 5, stripped[:120], ln=True)
            pdf.set_font("Helvetica", size=11)

    buf = BytesIO()
    try:
        pdf.output(buf)
    except Exception:
        pdf2 = FPDF()
        pdf2.add_page()
        pdf2.set_font("Helvetica", size=10)
        for line in text.replace('\n', ' ').split('. '):
            s = line.strip()
            if s:
                try:
                    pdf2.multi_cell(0, 6, s[:200] + '.')
                except Exception:
                    pdf2.cell(0, 6, s[:120], ln=True)
        pdf2.output(buf)
    buf.seek(0)
    return buf

def _sanitize(text):
    """Replace Unicode characters that break Helvetica font rendering."""
    # Character replacements for common Unicode punctuation
    replacements = {
        '\u2013': '-',   # en-dash
        '\u2014': '--',  # em-dash
        '\u2018': "'",   # left single quote
        '\u2019': "'",   # right single quote
        '\u201c': '"',   # left double quote
        '\u201d': '"',   # right double quote
        '\u2026': '...', # ellipsis
        '\u00a0': ' ',   # non-breaking space
        '\u2010': '-',   # hyphen
        '\u2011': '-',   # non-breaking hyphen
        '\u2012': '-',   # figure dash
        '\u2032': "'",   # prime
        '\u2033': "''",  # double prime
        '\u2212': '-',   # minus sign
        '\u00d7': 'x',   # multiplication sign
        '\u2022': '-',   # bullet
        '\u25e6': '-',   # white bullet
        '\u25aa': '-',   # black small square
        '\u25cf': '*',   # black circle
        '\u2192': '->',  # right arrow
        '\u2190': '<-',  # left arrow
        '\u00b7': '-',   # middle dot
        '\u00a9': '(c)', # copyright
        '\u00ae': '(R)', # registered
        '\u2122': '(TM)',# trademark
        '\u20ac': 'EUR', # euro
        '\u00a3': 'GBP', # pound
        '\u00a5': 'JPY', # yen
        '\u2264': '<=',  # less-or-equal
        '\u2265': '>=',  # greater-or-equal
        '\u2260': '!=',  # not equal
        '\u00b1': '+/-', # plus-minus
        '\u00f7': '/',   # division
        '\u2248': '~',   # almost equal
    }
    for unicode_char, ascii_replacement in replacements.items():
        text = text.replace(unicode_char, ascii_replacement)

    out = []
    for ch in text:
        code = ord(ch)
        if code < 32 and code not in (9, 10, 13):
            out.append(' ')
        elif 0x1F000 <= code <= 0x1FFFF:
            out.append('?')
        elif code > 0xFFFF:
            out.append('?')
        else:
            out.append(ch)
    return ''.join(out)

# ── Token Estimation ──
def estimate_tokens(text):
    return max(1, len(str(text)) // 4)

def estimate_cost(input_chars, output_tokens):
    in_tokens = estimate_tokens(input_chars)
    cost = (in_tokens / 1_000_000) * INPUT_COST_PER_1M
    cost += (output_tokens / 1_000_000) * OUTPUT_COST_PER_1M
    return in_tokens, cost

# ── Run Folder & Save ──
def get_job_folder(jd_text="", company_name=""):
    """Get or create a job-specific folder: CompanyName_YYYY-MM-DD. Cached per session."""
    try:
        existing = st.session_state.get("_job_folder")
        if existing and os.path.exists(existing):
            return existing
    except Exception:
        pass

    # Use company name if provided, else extract from JD first line
    if company_name and company_name.strip():
        folder_name = re.sub(r'[^a-zA-Z0-9_-]', '_', company_name.strip())[:40].strip('_')
    elif jd_text:
        first_line = jd_text.split('\n')[0].strip()
        for prefix in ['Job Description:', 'Job Title:', 'Position:', 'Role:', 'About the role', 'About this role']:
            if first_line.lower().startswith(prefix.lower()):
                first_line = first_line[len(prefix):].strip()
        folder_name = re.sub(r'[^a-zA-Z0-9_-]', '_', first_line)[:40].strip('_')
    else:
        folder_name = "Untitled_Job"

    ts = datetime.now().strftime("%Y-%m-%d")
    folder_name_full = f"{folder_name}_{ts}"
    base_dir = os.path.join(os.getcwd(), "outputs", folder_name_full)
    os.makedirs(base_dir, exist_ok=True)

    try:
        st.session_state["_job_folder"] = base_dir
    except Exception:
        pass

    return base_dir

def create_run_folder(label=""):
    """Create a named sub-folder inside the job-specific parent folder."""
    safe = re.sub(r'[^a-zA-Z0-9_-]', '_', label.strip())[:40] if label.strip() else "output"
    jd_text = ""
    company = ""
    try:
        jd_text = st.session_state.get("shared_jd", "")
        company = st.session_state.get("_company_name", "")
    except Exception:
        pass
    base_dir = get_job_folder(jd_text, company)
    run_dir = os.path.join(base_dir, safe)
    os.makedirs(run_dir, exist_ok=True)
    return run_dir

def save_run_file(content, filename, run_folder):
    filepath = os.path.join(run_folder, filename)
    if isinstance(content, BytesIO):
        with open(filepath, 'wb') as f:
            f.write(content.getvalue())
    else:
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(str(content))
    return filepath

def show_folder_summary(run_folder, saved_files):
    """Display saved files summary with job folder context."""
    rel = os.path.relpath(run_folder, os.getcwd())
    st.success(f":open_file_folder: `{rel}`")
    with st.expander(":page_facing_up: Saved files", expanded=True):
        for f in saved_files:
            st.caption(f"  {f}")

# ── File Naming Convention ──
def cl_filename(ext, job_title=""):
    """Cover_letter_{JobTitle}_{USERNAME}.{ext}"""
    from config import USERNAME, get_job_title
    jt = job_title if job_title else "Job"
    jt = re.sub(r'[^a-zA-Z0-9_-]', '_', jt)[:40].strip('_')
    return f"Cover_letter_{jt}_{USERNAME}.{ext}"

def cv_filename(ext, job_title=""):
    """CV_{JobTitle}_{USERNAME}.{ext}"""
    from config import USERNAME, get_job_title
    jt = job_title if job_title else "Job"
    jt = re.sub(r'[^a-zA-Z0-9_-]', '_', jt)[:40].strip('_')
    return f"CV_{jt}_{USERNAME}.{ext}"
